"""
Convertit MEMOIRE_BABIFIX_REDIGE.md en MEMOIRE_BABIFIX_REDIGE.docx
au format académique professionnel :
- Times New Roman 12, interligne 1,5, marges 2,5 cm partout
- Paragraphes justifiés avec retrait de première ligne
- Page de garde centrée verticalement, sans numéro de page
- Pages liminaires (Dédicace -> Abstract) : numérotation romaine (i, ii, iii)
- Corps (Introduction -> Annexes)            : numérotation arabe (1, 2, 3)
- Saut de page à chaque grande section
- Tableaux propres, en-tête grisé, bordures noires
- Styles Heading 1/2/3/4 natifs (table des matières automatique possible)
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

SRC = Path("MEMOIRE_BABIFIX_REDIGE.md")
DST = Path("MEMOIRE_BABIFIX_REDIGE.docx")

# Sections qui doivent démarrer sur une nouvelle page
LIMINARY_SECTIONS = {
    "DÉDICACE",
    "REMERCIEMENTS",
    "SOMMAIRE",
    "AVANT-PROPOS",
    "LISTE DES SIGLES ET ABRÉVIATIONS",
    "LISTE DES TABLEAUX",
    "LISTE DES FIGURES",
    "LISTE DES ANNEXES",
    "RÉSUMÉ",
    "ABSTRACT",
}

# Titre déclencheur du passage de numérotation romaine -> arabe
INTRO_TRIGGER = "INTRODUCTION"

# Sections principales (## ou #) qui démarrent une nouvelle page
MAJOR_BREAK_SECTIONS = LIMINARY_SECTIONS | {
    "INTRODUCTION",
    "CONCLUSION",
    "RÉFÉRENCES BIBLIOGRAPHIQUES",
    "WEBOGRAPHIES",
    "ANNEXES",
}


# ============================================================================
# UTILITAIRES XML
# ============================================================================
def set_cell_borders(cell, color: str = "000000", size: str = "6") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill_hex: str = "D9E2F3") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_section_page_numbering(section, fmt: str = "decimal", start: int | None = None) -> None:
    """
    fmt : "decimal" (1,2,3) ou "lowerRoman" (i,ii,iii) ou "upperRoman"
    start : valeur de départ (None = continuer)
    """
    sectPr = section._sectPr
    # Retirer ancien pgNumType
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def add_page_number_to_footer(section, suppress: bool = False) -> None:
    """Ajoute un numéro de page centré au pied de page (ou rien si suppress)."""
    footer = section.footer
    # Vide le footer
    for p in list(footer.paragraphs):
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if suppress:
        return

    # Insertion du champ PAGE
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE  \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    # Police du footer
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def set_vertical_alignment_center(section) -> None:
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:vAlign")):
        sectPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    sectPr.append(vAlign)


# ============================================================================
# STYLES DOCUMENT
# ============================================================================
def setup_doc_styles(doc: Document) -> None:
    """Mise en forme par défaut + styles Heading personnalisés."""
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Heading 1 : grand titre (Partie, Chapitre)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x0B, 0x2A, 0x4A)
    pf = h1.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True

    # Heading 2 : grande section (DÉDICACE, REMERCIEMENTS, etc.)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x0B, 0x2A, 0x4A)
    pf = h2.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(18)
    pf.space_after = Pt(12)
    pf.keep_with_next = True

    # Heading 3 : section numérotée (1.1, 2.3, etc.)
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x14, 0x3C, 0x6E)
    pf = h3.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

    # Heading 4 : sous-section
    h4 = doc.styles["Heading 4"]
    h4.font.name = "Times New Roman"
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0x14, 0x3C, 0x6E)
    pf = h4.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.keep_with_next = True


def setup_margins(section) -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ============================================================================
# CONTENU PARAGRAPHE (gras / italique)
# ============================================================================
def add_inline_runs(paragraph, text: str, base_bold: bool = False,
                    base_italic: bool = False, base_size: int = 12) -> None:
    """Gère **gras** et *italique* dans une chaîne de texte."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = "Times New Roman"
            run.font.size = Pt(base_size)
            run.bold = base_bold
            run.italic = base_italic
        token = m.group()
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.font.name = "Times New Roman"
            run.font.size = Pt(base_size)
            run.bold = True
            run.italic = base_italic
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Times New Roman"
            run.font.size = Pt(base_size)
            run.bold = base_bold
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Times New Roman"
        run.font.size = Pt(base_size)
        run.bold = base_bold
        run.italic = base_italic


# ============================================================================
# TABLEAUX
# ============================================================================
def add_table_from_md(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row in enumerate(rows):
        is_header = (i == 0)
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if not is_header else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, val.strip(), base_bold=is_header, base_size=11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            if is_header:
                shade_cell(cell, fill_hex="D9E2F3")

    # Espace après le tableau
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


# ============================================================================
# AJOUT DE PARAGRAPHES
# ============================================================================
def add_paragraph_justified(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    add_inline_runs(p, text)


def add_centered_paragraph(doc: Document, text: str, size: int = 12,
                           bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, text, base_bold=bold, base_italic=italic, base_size=size)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_break(doc: Document, start_type=WD_SECTION_START.NEW_PAGE):
    """Insère un saut de section (Next Page) et renvoie la nouvelle section."""
    new_section = doc.add_section(start_type)
    setup_margins(new_section)
    return new_section


# ============================================================================
# PAGE DE GARDE (composition manuelle, centrée)
# ============================================================================
def build_title_page(doc: Document, lines: list[str]) -> None:
    """
    Compose la page de garde de manière soignée :
    bandeau République / Ministère / IIT en haut, titre au centre, auteur + maître de stage en bas.
    `lines` est l'ensemble brut des lignes Markdown de la page de garde.
    """
    # On rend la première section : page de garde
    # Lignes manuelles pour un rendu maîtrisé
    add_centered_paragraph(doc, "République de Côte d'Ivoire", size=14, bold=True)
    add_centered_paragraph(doc, "Union — Discipline — Travail", size=11, italic=True)
    add_centered_paragraph(doc, "Ministère de l'Enseignement Supérieur et de la Recherche Scientifique",
                           size=11, bold=False)

    # Espace
    for _ in range(2):
        doc.add_paragraph()

    add_centered_paragraph(doc, "INSTITUT IVOIRIEN DE TECHNOLOGIE", size=16, bold=True)
    add_centered_paragraph(doc, "DÉPARTEMENT D'INFORMATIQUE", size=13, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    add_centered_paragraph(doc, "MÉMOIRE DE FIN DE CYCLE", size=18, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    add_centered_paragraph(doc, "THÈME :", size=12, bold=True)
    add_centered_paragraph(doc,
                           "Conception et réalisation d'une plateforme numérique de services à domicile en Côte d'Ivoire — le cas BABIFIX",
                           size=15, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    # Bas de page : auteur, encadrants, période
    add_centered_paragraph(doc, "Réalisé par :  [NOM Prénom]", size=12, bold=True)
    add_centered_paragraph(doc, "Élève en [Cycle / Filière], Promotion [année]", size=11, italic=True)

    doc.add_paragraph()

    add_centered_paragraph(doc, "Directeur de mémoire :  [Nom, Prénom et titre]", size=12)
    add_centered_paragraph(doc,
                           "Maître de stage :  M. TIÉ Narcisse, Ingénieur Informaticien",
                           size=12)
    add_centered_paragraph(doc, "Société Nationale de Développement Informatique (SNDI)",
                           size=11, italic=True)

    doc.add_paragraph()
    add_centered_paragraph(doc, "Année académique :  [année]", size=12, bold=True)


# ============================================================================
# PARSER PRINCIPAL
# ============================================================================
def parse_markdown(md_text: str) -> Document:
    doc = Document()
    setup_doc_styles(doc)
    setup_margins(doc.sections[0])

    # SECTION 1 : Page de garde -- pas de numéro
    first_section = doc.sections[0]
    set_vertical_alignment_center(first_section)
    add_page_number_to_footer(first_section, suppress=True)

    # Compose la page de garde manuellement (on ignore les lignes brutes du MD)
    build_title_page(doc, [])

    # On passe en mode parsing normal pour le reste
    # SECTION 2 : pages liminaires (Dédicace -> Abstract) en romain
    sec2 = add_section_break(doc, WD_SECTION_START.NEW_PAGE)
    set_section_page_numbering(sec2, fmt="lowerRoman", start=1)
    add_page_number_to_footer(sec2)

    # Drapeau pour savoir si on a déjà basculé en numérotation arabe
    arabic_started = False

    # Découpe le markdown en lignes, on ignore l'entête page de garde
    lines = md_text.split("\n")
    i = 0
    # Saute toute la page de garde et atteint la première ##
    while i < len(lines) and not lines[i].strip().startswith("## "):
        i += 1

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Lignes vides
        if not line.strip():
            i += 1
            continue

        # Séparateurs ---
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # === TITRES ===
        # H1 : grandes parties (PREMIÈRE PARTIE / CHAPITRE I, etc.)
        if line.startswith("# "):
            title = line[2:].strip()
            # Saut de page avant un H1
            doc.add_page_break()
            p = doc.add_paragraph(title, style="Heading 1")
            i += 1
            continue

        # H2 : grandes sections (DÉDICACE, INTRODUCTION, CONCLUSION, etc.)
        if line.startswith("## "):
            title = line[3:].strip()
            normalized = title.upper().split("*")[0].strip()
            # On retire les marqueurs *Page X*
            normalized = re.sub(r"\s*\(.*\)\s*$", "", normalized)

            # Saut de page si nécessaire
            if normalized in MAJOR_BREAK_SECTIONS or any(
                normalized.startswith(s) for s in MAJOR_BREAK_SECTIONS
            ):
                # Si on arrive à l'INTRODUCTION, on bascule en numérotation arabe
                if not arabic_started and (normalized == INTRO_TRIGGER or normalized.startswith(INTRO_TRIGGER)):
                    sec3 = add_section_break(doc, WD_SECTION_START.NEW_PAGE)
                    set_section_page_numbering(sec3, fmt="decimal", start=1)
                    add_page_number_to_footer(sec3)
                    arabic_started = True
                else:
                    doc.add_page_break()

            p = doc.add_paragraph(title, style="Heading 2")
            i += 1
            continue

        # H3 et H4 (sous-sections numérotées)
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(line[5:].strip(), style="Heading 4")
            i += 1
            continue

        # === TABLEAUX ===
        if line.lstrip().startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            rows = []
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(header_cells)
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row_cells)
                i += 1
            add_table_from_md(doc, rows)
            continue

        # === LISTES À PUCES ===
        if line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            add_bullet(doc, line.lstrip()[2:])
            i += 1
            continue

        # === CITATIONS > ===
        if line.lstrip().startswith("> "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Cm(2)
            p.paragraph_format.right_indent = Cm(2)
            add_inline_runs(p, line.lstrip()[2:], base_italic=True)
            i += 1
            continue

        # === PARAGRAPHE NORMAL ===
        # On regroupe les lignes consécutives en un seul paragraphe
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].rstrip()
            if not next_line.strip():
                break
            if next_line.startswith(("#", "|", "-", "*", "> ", "---")):
                break
            para_lines.append(next_line)
            j += 1
        full_text = " ".join(para_lines)
        # Légende de tableau / figure : on la met en gras centrée sans retrait
        if full_text.startswith("**Tableau ") or full_text.startswith("**Figure "):
            # Retire les ** englobants
            clean = full_text.strip()
            add_centered_paragraph(doc, clean.strip("*").strip(), size=11, bold=True, italic=True)
        else:
            # Paragraphes centrés pour la dédicace et la zone résumé/abstract du début
            first_line_indent = True
            add_paragraph_justified(doc, full_text, first_line_indent=first_line_indent)
        i = j

    return doc


# ============================================================================
# MAIN
# ============================================================================
def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = parse_markdown(md)
    out = DST
    try:
        doc.save(str(out))
    except PermissionError:
        # Si le fichier est ouvert dans Word, on écrit une version horodatée
        from datetime import datetime
        out = DST.with_name(f"MEMOIRE_BABIFIX_REDIGE_{datetime.now():%H%M%S}.docx")
        doc.save(str(out))
        print(f"[!] Le fichier {DST.name} est ouvert dans Word.")
    print(f"OK : {out.resolve()}")


if __name__ == "__main__":
    main()
